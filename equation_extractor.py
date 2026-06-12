#!/usr/bin/env python3
"""
Equation Extractor and Compiler
Extracts mathematical, scientific, physics, chemistry, and biology equations
from various document formats and images, compiling them into a Word document
using ASCIIMath notation.
"""

import os
import re
import sys
import json
import argparse
import hashlib
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Set, Any
import logging
import sqlite3
from dataclasses import dataclass, asdict
from enum import Enum

# Third-party imports
try:
    import PyPDF2
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pythoncom
    from win32com import client as win32client
    import fitz  # PyMuPDF
    import cv2
    import numpy as np
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.cluster import KMeans
    import matplotlib.pyplot as plt
    import seaborn as sns
except ImportError as e:
    print(f"Missing required library: {e}")
    print("Please install required packages:")
    print("pip install PyPDF2 pdf2image pytesseract python-docx pywin32 PyMuPDF opencv-python scikit-learn matplotlib seaborn")
    print("Also install:")
    print("1. Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")
    print("2. Poppler for PDF to image conversion: https://github.com/oschwartz10612/poppler-windows/releases/")
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('equation_extractor.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class EquationType(Enum):
    """Enum for equation types."""
    MATHEMATICAL = "mathematical"
    PHYSICS = "physics"
    CHEMISTRY = "chemistry"
    BIOLOGY = "biology"
    UNCATEGORIZED = "uncategorized"

@dataclass
class Equation:
    """Data class for storing equation information."""
    id: str
    content: str
    asciimath: str
    equation_type: EquationType
    source_file: str
    page_number: Optional[int] = None
    confidence: float = 0.0
    extracted_date: str = ""
    tags: List[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.tags is None:
            self.tags = []
        if self.metadata is None:
            self.metadata = {}
        if not self.extracted_date:
            self.extracted_date = datetime.now().isoformat()

class EquationDatabase:
    """SQLite database for storing equations."""
    
    def __init__(self, db_path: str = "equations.db"):
        self.db_path = db_path
        self._init_database()
    
    def _init_database(self):
        """Initialize the database schema."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Create equations table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS equations (
                id TEXT PRIMARY KEY,
                content TEXT NOT NULL,
                asciimath TEXT NOT NULL,
                equation_type TEXT NOT NULL,
                source_file TEXT NOT NULL,
                page_number INTEGER,
                confidence REAL DEFAULT 0.0,
                extracted_date TEXT,
                tags TEXT,
                metadata TEXT
            )
        ''')
        
        # Create index for faster searching
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_equation_type 
            ON equations(equation_type)
        ''')
        
        cursor.execute('''
            CREATE INDEX IF NOT EXISTS idx_source_file 
            ON equations(source_file)
        ''')
        
        conn.commit()
        conn.close()
    
    def save_equation(self, equation: Equation):
        """Save an equation to the database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT OR REPLACE INTO equations 
            (id, content, asciimath, equation_type, source_file, 
             page_number, confidence, extracted_date, tags, metadata)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            equation.id,
            equation.content,
            equation.asciimath,
            equation.equation_type.value,
            equation.source_file,
            equation.page_number,
            equation.confidence,
            equation.extracted_date,
            json.dumps(equation.tags),
            json.dumps(equation.metadata)
        ))
        
        conn.commit()
        conn.close()
    
    def get_equations_by_type(self, equation_type: EquationType) -> List[Equation]:
        """Get all equations of a specific type."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM equations WHERE equation_type = ?
        ''', (equation_type.value,))
        
        rows = cursor.fetchall()
        conn.close()
        
        equations = []
        for row in rows:
            equations.append(Equation(
                id=row[0],
                content=row[1],
                asciimath=row[2],
                equation_type=EquationType(row[3]),
                source_file=row[4],
                page_number=row[5],
                confidence=row[6],
                extracted_date=row[7],
                tags=json.loads(row[8]) if row[8] else [],
                metadata=json.loads(row[9]) if row[9] else {}
            ))
        
        return equations
    
    def get_all_equations(self) -> List[Equation]:
        """Get all equations from the database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM equations')
        rows = cursor.fetchall()
        conn.close()
        
        equations = []
        for row in rows:
            equations.append(Equation(
                id=row[0],
                content=row[1],
                asciimath=row[2],
                equation_type=EquationType(row[3]),
                source_file=row[4],
                page_number=row[5],
                confidence=row[6],
                extracted_date=row[7],
                tags=json.loads(row[8]) if row[8] else [],
                metadata=json.loads(row[9]) if row[9] else {}
            ))
        
        return equations
    
    def delete_equation(self, equation_id: str):
        """Delete an equation from the database."""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('DELETE FROM equations WHERE id = ?', (equation_id,))
        conn.commit()
        conn.close()

class EquationExtractor:
    """Main class for extracting equations from various formats."""
    
    # Enhanced equation patterns
    EQUATION_PATTERNS = {
        EquationType.MATHEMATICAL: [
            # LaTeX patterns
            r'\$.*?\$',
            r'\\\(.*?\\\)',
            r'\\\[.*?\\\]',
            r'\\begin\{equation\}.*?\\end\{equation\}',
            r'\\begin\{align\}.*?\\end\{align\}',
            r'\\begin\{gather\}.*?\\end\{gather\}',
            r'\\begin\{multline\}.*?\\end\{multline\}',
            
            # Mathematical symbols
            r'[∑∫∏∮∞√∛∜±∓×÷∙∘•∗⊕⊗⊙⊚⊛⊞⊟⋇⋈⋉⋊⋋⋌≅≈≃≂≀≁≄≆≇≉≊≋≌≍≎≏≐≑≒≓≔≕≖≗≘≙≚≛≜≝≞≟≠≡≢≣≤≥≦≧≨≩≪≫≬≭≮≯≰≱≲≳≴≵≶≷≸≹≺≻≼≽≾≿⊀⊁⊂⊃⊄⊅⊆⊇⊈⊉⊊⊋⊌⊍⊎⊏⊐⊑⊒⊓⊔⊕⊖⊗⊘⊙⊚⊛⊜⊝⊞⊟⊠⊡⊢⊣⊤⊥⊦⊧⊨⊩⊪⊫⊬⊭⊮⊯⊰⊱⊲⊳⊴⊵⊶⊷⊸⊹⊺⊻⊼⊽⊾⊿⋀⋁⋂⋃⋄⋅⋆⋇⋈⋉⋊⋋⋌⋍⋎⋏⋐⋑⋒⋓⋔⋕⋖⋗⋘⋙⋚⋛⋜⋝⋞⋟⋠⋡⋢⋣⋤⋥⋦⋧⋨⋩⋪⋫⋬⋭⋮⋯⋰⋱]',
            
            # Functions
            r'\\sin|\\cos|\\tan|\\log|\\ln|\\exp|\\lim|\\sum|\\prod|\\int',
            
            # Fractions and roots
            r'\\frac\{.*?\}\{.*?\}',
            r'\\sqrt\{.*?\}',
            r'\\root\{.*?\}\{.*?\}',
            
            # Superscripts and subscripts
            r'\^\{.*?\}',
            r'_\s*\{.*?\}',
        ],
        EquationType.PHYSICS: [
            # Newton's laws
            r'F\s*=\s*ma',
            r'F\s*=\s*G\frac\{m_1m_2\}\{r\^2\}',
            
            # Energy equations
            r'E\s*=\s*mc\^2',
            r'E\s*=\s*hf',
            r'E\s*=\s*\frac\{1\}\{2\}mv\^2',
            r'PE\s*=\s*mgh',
            
            # Kinematics
            r'v\s*=\s*u\s*\+\s*at',
            r's\s*=\s*ut\s*\+\s*\frac\{1\}\{2\}at\^2',
            r'v\^2\s*=\s*u\^2\s*\+\s*2as',
            
            # Electricity
            r'V\s*=\s*IR',
            r'P\s*=\s*VI',
            r'P\s*=\s*I\^2R',
            
            # Optics
            r'\frac\{1\}\{f\}\s*=\s*\frac\{1\}\{u\}\s*\+\s*\frac\{1\}\{v\}',
            r'n\s*=\s*\frac\{c\}\{v\}',
            
            # Thermodynamics
            r'PV\s*=\s*nRT',
            r'Q\s*=\s*mcΔT',
        ],
        EquationType.CHEMISTRY: [
            # Chemical reactions
            r'[A-Z][a-z]?\s*\+\s*[A-Z][a-z]?\s*[→⇌⇋⇄↔]\s*[A-Z][a-z]?',
            r'[A-Z][a-z]?\d?\s*\+\s*[A-Z][a-z]?\d?\s*[→⇌⇋⇄↔]\s*[A-Z][a-z]?\d?',
            
            # Common compounds
            r'H_2O',
            r'CO_2',
            r'CH_4',
            r'C_6H_\{12\}O_6',
            r'NaOH\s*\+\s*HCl\s*→\s*NaCl\s*\+\s*H_2O',
            
            # Gas laws
            r'PV\s*=\s*nRT',
            r'P_1V_1\s*=\s*P_2V_2',
            
            # Equilibrium
            r'K\s*=\s*\frac\{\[C\]\^c\[D\]\^d\}\{\[A\]\^a\[B\]\^b\}',
            
            # Thermodynamics
            r'ΔG\s*=\s*ΔH\s*-\s*TΔS',
            r'ΔG°\s*=\s*-RT\\ln\ K',
            
            # Quantum chemistry
            r'E\s*=\s*hν',
            r'λ\s*=\s*\frac\{h\}\{p\}',
            
            # Radioactivity
            r'N\s*=\s*N_0e\^\{-λt\}',
            r't_\{\frac\{1\}\{2\}\}\s*=\s*\frac\{\\ln\ 2\}\{λ\}',
        ],
        EquationType.BIOLOGY: [
            # Photosynthesis and respiration
            r'6CO_2\s*\+\s*6H_2O\s*→\s*C_6H_\{12\}O_6\s*\+\s*6O_2',
            r'C_6H_\{12\}O_6\s*\+\s*6O_2\s*→\s*6CO_2\s*\+\s*6H_2O\s*\+\s*energy',
            
            # Central dogma
            r'DNA\s*→\s*RNA\s*→\s*Protein',
            
            # Genetics
            r'p\^2\s*\+\s*2pq\s*\+\s*q\^2\s*=\s*1',
            
            # Population growth
            r'N_t\s*=\s*N_0\s*\*\s*2\^t',
            r'\frac\{dN\}\{dt\}\s*=\s*rN\(1\s*-\s*\frac\{N\}\{K\}\)',
            
            # Enzyme kinetics
            r'v\s*=\s*\frac\{V_\{max\}\[S\]\}\{K_m\s*\+\s*\[S\]\}',
            
            # pH calculations
            r'pH\s*=\s*-\\log\[H\^+\]',
            r'pOH\s*=\s*-\\log\[OH\^-\]',
            
            # Hardy-Weinberg
            r'p\s*\+\s*q\s*=\s*1',
        ]
    }
    
    def __init__(self):
        self.db = EquationDatabase()
        self.image_formats = {'.jpg', '.jpeg', '.png', '.gif', '.tiff', '.tif', '.bmp', '.webp'}
        self.document_formats = {'.pdf', '.txt', '.doc', '.docx', '.odt', '.rtf'}
        self._init_tesseract()
    
    def _init_tesseract(self):
        """Initialize Tesseract OCR paths."""
        # Common Tesseract installation paths
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            '/usr/bin/tesseract',
            '/usr/local/bin/tesseract'
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                logger.info(f"Tesseract found at: {path}")
                return
        
        logger.warning("Tesseract not found. OCR functionality may not work.")
    
    def extract_from_file(self, file_path: str) -> List[Equation]:
        """Extract equations from any supported file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return []
        
        equations = []
        
        try:
            if file_path.suffix.lower() in self.image_formats:
                equations = self._extract_from_image(str(file_path))
            elif file_path.suffix.lower() == '.pdf':
                equations = self._extract_from_pdf(str(file_path))
            elif file_path.suffix.lower() == '.txt':
                equations = self._extract_from_txt(str(file_path))
            elif file_path.suffix.lower() in {'.docx', '.doc'}:
                equations = self._extract_from_docx(str(file_path))
            elif file_path.suffix.lower() in {'.odt', '.odf'}:
                equations = self._extract_from_odf(str(file_path))
            else:
                logger.warning(f"Unsupported file format: {file_path.suffix}")
            
            # Save extracted equations to database
            for equation in equations:
                self.db.save_equation(equation)
            
            logger.info(f"Extracted {len(equations)} equations from {file_path.name}")
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
        
        return equations
    
    def _extract_from_image(self, image_path: str) -> List[Equation]:
        """Extract equations from image files using OCR."""
        equations = []
        
        try:
            # Preprocess image for better OCR
            image = Image.open(image_path)
            
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Apply sharpening
            image = image.filter(ImageFilter.SHARPEN)
            
                        # Use Tesseract with custom configuration for equations
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist="0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ+-=()[]{}^_/\\|<>.,:;!?@#$%&*αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ∑∫∏√∞≈≠≤≥→⇌⇋"'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            # Extract equations from OCR text
            extracted = self._extract_equations_from_text(text, source_file=image_path)
            
            for eq_type, eq_content in extracted:
                # Generate ASCIIMath representation
                asciimath = self._convert_to_asciimath(eq_content)
                
                # Create equation object
                equation = Equation(
                    id=self._generate_id(eq_content, image_path),
                    content=eq_content,
                    asciimath=asciimath,
                    equation_type=eq_type,
                    source_file=image_path,
                    confidence=0.8,  # OCR confidence estimate
                    tags=["image_ocr", "extracted"]
                )
                equations.append(equation)
                
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {e}")
            
        return equations
    
    def _extract_from_pdf(self, pdf_path: str) -> List[Equation]:
        """Extract equations from PDF files."""
        equations = []
        
        try:
            # Method 1: Extract text directly from PDF
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content = page.extract_text()
                    if text_content:
                        extracted = self._extract_equations_from_text(
                            text_content, 
                            source_file=pdf_path,
                            page_number=page_num + 1
                        )
                        
                        for eq_type, eq_content in extracted:
                            asciimath = self._convert_to_asciimath(eq_content)
                            equation = Equation(
                                id=self._generate_id(eq_content, pdf_path, page_num),
                                content=eq_content,
                                asciimath=asciimath,
                                equation_type=eq_type,
                                source_file=pdf_path,
                                page_number=page_num + 1,
                                confidence=0.9,
                                tags=["pdf_text", "direct_extraction"]
                            )
                            equations.append(equation)
            
            # Method 2: Use OCR for scanned PDFs
            try:
                images = convert_from_path(pdf_path, dpi=300)
                for page_num, image in enumerate(images):
                    # Convert to grayscale and enhance
                    if image.mode != 'L':
                        image = image.convert('L')
                    
                    enhancer = ImageEnhance.Contrast(image)
                    image = enhancer.enhance(2.0)
                    
                    # OCR with equation-specific config
                    custom_config = r'--oem 3 --psm 6'
                    ocr_text = pytesseract.image_to_string(image, config=custom_config)
                    
                    extracted = self._extract_equations_from_text(
                        ocr_text,
                        source_file=pdf_path,
                        page_number=page_num + 1
                    )
                    
                    for eq_type, eq_content in extracted:
                        asciimath = self._convert_to_asciimath(eq_content)
                        equation = Equation(
                            id=self._generate_id(eq_content, pdf_path, page_num, "ocr"),
                            content=eq_content,
                            asciimath=asciimath,
                            equation_type=eq_type,
                            source_file=pdf_path,
                            page_number=page_num + 1,
                            confidence=0.7,
                            tags=["pdf_ocr", "scanned"]
                        )
                        equations.append(equation)
                        
            except Exception as ocr_error:
                logger.warning(f"PDF OCR failed for {pdf_path}: {ocr_error}")
                
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            
        return equations
    
    def _extract_from_txt(self, txt_path: str) -> List[Equation]:
        """Extract equations from text files."""
        equations = []
        
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                extracted = self._extract_equations_from_text(content, source_file=txt_path)
                
                for eq_type, eq_content in extracted:
                    asciimath = self._convert_to_asciimath(eq_content)
                    equation = Equation(
                        id=self._generate_id(eq_content, txt_path),
                        content=eq_content,
                        asciimath=asciimath,
                        equation_type=eq_type,
                        source_file=txt_path,
                        confidence=1.0,
                        tags=["text_file", "plain_text"]
                    )
                    equations.append(equation)
                    
        except Exception as e:
            logger.error(f"Error processing text file {txt_path}: {e}")
            
        return equations
    
    def _extract_from_docx(self, docx_path: str) -> List[Equation]:
        """Extract equations from DOCX files."""
        equations = []
        
        try:
            doc = docx.Document(docx_path)
            
            # Extract from paragraphs
            for para_num, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    extracted = self._extract_equations_from_text(
                        para.text,
                        source_file=docx_path,
                        paragraph_num=para_num
                    )
                    
                    for eq_type, eq_content in extracted:
                        asciimath = self._convert_to_asciimath(eq_content)
                        equation = Equation(
                            id=self._generate_id(eq_content, docx_path, para_num),
                            content=eq_content,
                            asciimath=asciimath,
                            equation_type=eq_type,
                            source_file=docx_path,
                            confidence=0.95,
                            tags=["docx", "paragraph"]
                        )
                        equations.append(equation)
            
            # Extract from tables
            for table_num, table in enumerate(doc.tables):
                for row_num, row in enumerate(table.rows):
                    for cell_num, cell in enumerate(row.cells):
                        if cell.text.strip():
                            extracted = self._extract_equations_from_text(
                                cell.text,
                                source_file=docx_path,
                                table_num=table_num,
                                row_num=row_num,
                                cell_num=cell_num
                            )
                            
                            for eq_type, eq_content in extracted:
                                asciimath = self._convert_to_asciimath(eq_content)
                                equation = Equation(
                                    id=self._generate_id(eq_content, docx_path, table_num, row_num, cell_num),
                                    content=eq_content,
                                    asciimath=asciimath,
                                    equation_type=eq_type,
                                    source_file=docx_path,
                                    confidence=0.95,
                                    tags=["docx", "table"]
                                )
                                equations.append(equation)
                                
        except Exception as e:
            logger.error(f"Error processing DOCX file {docx_path}: {e}")
            
        return equations
    
    def _extract_from_odf(self, odf_path: str) -> List[Equation]:
        """Extract equations from ODF files."""
        equations = []
        
        try:
            # For ODF files, we'll use text extraction
            # Note: This is a simplified approach
            with open(odf_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                extracted = self._extract_equations_from_text(content, source_file=odf_path)
                
                for eq_type, eq_content in extracted:
                    asciimath = self._convert_to_asciimath(eq_content)
                    equation = Equation(
                        id=self._generate_id(eq_content, odf_path),
                        content=eq_content,
                        asciimath=asciimath,
                        equation_type=eq_type,
                        source_file=odf_path,
                        confidence=0.85,
                        tags=["odf", "open_document"]
                    )
                    equations.append(equation)
                    
        except Exception as e:
            logger.error(f"Error processing ODF file {odf_path}: {e}")
            
        return equations
    
    def _extract_equations_from_text(self, text: str, **kwargs) -> List[Tuple[EquationType, str]]:
        """Extract equations from text and classify them."""
        equations = []
        
        # Split text into lines for better processing
        lines = text.split('\n')
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for each equation type
            for eq_type, patterns in self.EQUATION_PATTERNS.items():
                for pattern in patterns:
                    try:
                        matches = re.findall(pattern, line, re.IGNORECASE | re.DOTALL)
                        for match in matches:
                            if self._is_valid_equation(match):
                                equations.append((eq_type, match))
                    except re.error:
                        continue
            
            # Also look for standalone equations (lines that look like equations)
            if self._looks_like_equation(line):
                # Try to classify it
                eq_type = self._classify_equation(line)
                equations.append((eq_type, line))
        
        # Remove duplicates while preserving order
        seen = set()
        unique_equations = []
        for eq_type, eq_content in equations:
            normalized = self._normalize_equation(eq_content)
            if normalized not in seen:
                seen.add(normalized)
                unique_equations.append((eq_type, eq_content))
        
        return unique_equations
    
    def _is_valid_equation(self, text: str) -> bool:
        """Check if text is a valid equation."""
        # Must contain at least one mathematical operator or symbol
        math_symbols = r'[=+\-*/^_()\[\]{}<>|\\∑∫∏√∞≈≠≤≥→]'
        if not re.search(math_symbols, text):
            return False
        
        # Must be reasonable length
        if len(text) < 2 or len(text) > 500:
            return False
        
        return True
    
    def _looks_like_equation(self, text: str) -> bool:
        """Heuristic check if text looks like an equation."""
        # Contains equals sign with something on both sides
        if '=' in text:
            parts = text.split('=')
            if len(parts) == 2 and len(parts[0].strip()) > 0 and len(parts[1].strip()) > 0:
                return True
        
        # Contains common mathematical patterns
        patterns = [
            r'^[A-Za-z]+\s*=\s*',
            r'\\[a-zA-Z]+\{',
            r'\^\{',
            r'_\s*\{',
            r'frac\{',
            r'sqrt\{',
            r'sum_\{',
            r'int_\{',
        ]
        
        for pattern in patterns:
            if re.search(pattern, text):
                return True
        
        return False
    
    def _classify_equation(self, equation: str) -> EquationType:
        """Classify an equation into a specific type."""
        equation_lower = equation.lower()
        
        # Check for physics patterns
        physics_keywords = ['f=', 'e=', 'v=', 'p=', 'ke=', 'pe=', 'w=', 'ρ=', 'λ=', 'n=']
        for keyword in physics_keywords:
            if keyword in equation_lower:
                return EquationType.PHYSICS
        
        # Check for chemistry patterns
        chem_keywords = ['h2o', 'co2', 'ch4', 'naoh', 'hcl', 'nacl', '→', '⇌', 'k=', 'δg', 'ph=']
        for keyword in chem_keywords:
            if keyword in equation_lower:
                return EquationType.CHEMISTRY
        
        # Check for biology patterns
        bio_keywords = ['c6h12o6', 'dna', 'rna', 'protein', 'vmax', 'km', 'nt=', 'n0']
        for keyword in bio_keywords:
            if keyword in equation_lower:
                return EquationType.BIOLOGY
        
        # Default to mathematical
        return EquationType.MATHEMATICAL
    
    def _normalize_equation(self, equation: str) -> str:
        """Normalize equation for duplicate detection."""
        # Remove extra whitespace
        normalized = re.sub(r'\s+', ' ', equation.strip())
        
        # Remove LaTeX command backslashes for comparison
        normalized = normalized.replace('\\', '')
        
        # Convert to lowercase
        normalized = normalized.lower()
        
        return normalized
    
    def _convert_to_asciimath(self, equation: str) -> str:
        """Convert equation to ASCIIMath notation."""
        # Basic LaTeX to ASCIIMath conversions
        conversions = [
            (r'\\frac\{(.*?)\}\{(.*?)\}', r'(\1)/(\2)'),
            (r'\\sqrt\{(.*?)\}', r'sqrt(\1)'),
            (r'\\sum_\{.*?\}', r'sum'),
            (r'\\int_\{.*?\}', r'int'),
            (r'\\lim_\{.*?\}', r'lim'),
            (r'\\sin', r'sin'),
            (r'\\cos', r'cos'),
            (r'\\tan', r'tan'),
            (r'\\log', r'log'),
            (r'\\ln', r'ln'),
            (r'\\exp', r'exp'),
            (r'\^\{', r'^('),
            (r'_\s*\{', r'_('),
            (r'\\[a-zA-Z]+\s*', ''),  # Remove other LaTeX commands
            (r'\$', ''),  # Remove dollar signs
            (r'\\\(', ''),  # Remove LaTeX inline markers
            (r'\\\)', ''),
            (r'\\\[', ''),  # Remove LaTeX display markers
            (r'\\\]', ''),
        ]
        
        asciimath = equation
        for pattern, replacement in conversions:
            asciimath = re.sub(pattern, replacement, asciimath)
        
        # Clean up extra spaces and braces
        asciimath = asciimath.replace('{', '(').replace('}', ')')
        asciimath = re.sub(r'\s+', ' ', asciimath.strip())
        
        return asciimath
    
    def _generate_id(self, *args) -> str:
        """Generate a unique ID for an equation."""
        input_string = '_'.join(str(arg) for arg in args)
        return hashlib.md5(input_string.encode()).hexdigest()[:16]
    
    def extract_from_directory(self, directory_path: str) -> List[Equation]:
        """Extract equations from all supported files in a directory."""
        directory = Path(directory_path)
        all_equations = []
        
        if not directory.exists():
            logger.error(f"Directory not found: {directory}")
            return []
        
        # Get all supported files
        supported_extensions = self.image_formats.union(self.document_formats)
        files = []
        for ext in supported_extensions:
            files.extend(directory.glob(f'*{ext}'))
            files.extend(directory.glob(f'*{ext.upper()}'))
        
        logger.info(f"Found {len(files)} supported files in {directory}")
        
        # Process each file
        for file_path in files:
            equations = self.extract_from_file(str(file_path))
            all_equations.extend(equations)
        
        return all_equations

class EquationCompiler:
    """Compile equations into Word document."""
    
    def __init__(self, output_path: str):
        self.output_path = Path(output_path)
        self.db = EquationDatabase()
    
    def compile_to_word(self, equations: List[Equation] = None, 
                       append: bool = True, 
                       sort_by_type: bool = True,
                       include_statistics: bool = True):
        """Compile equations to Word document."""
        try:
            # Get equations if not provided
            if equations is None:
                equations = self.db.get_all_equations()
            
            if not equations:
                logger.warning("No equations to compile")
                return False
            
            # Create or load document
            if append and self.output_path.exists():
                doc = docx.Document(self.output_path)
                # Add a page break before new content
                doc.add_page_break()
            else:
                doc = docx.Document()
            
            # Add title
            title = doc.add_heading('Equation Compendium', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp.runs[0].font.italic = True
            
            # Add statistics if requested
            if include_statistics:
                self._add_statistics(doc, equations)
            
            # Sort equations by type if requested
            if sort_by_type:
                equations_by_type = {}
                for eq_type in EquationType:
                    equations_by_type[eq_type] = [
                        eq for eq in equations if eq.equation_type == eq_type
                    ]
                
                # Add equations by type
                for eq_type, type_equations in equations_by_type.items():
                    if type_equations:
                        self._add_equation_section(doc, eq_type, type_equations)
            else:
                # Add all equations in one section
                self._add_equation_section(doc, "All Equations", equations)
            
            # Save document
            doc.save(self.output_path)
            logger.info(f"Compiled {len(equations)} equations to {self.output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error compiling to Word: {e}")
            return False
    
    def _add_statistics(self, doc, equations: List[Equation]):
        """Add statistics section to document."""
             stats_heading = doc.add_heading('Statistics', 1)
        
        # Count equations by type
        type_counts = {}
        for eq_type in EquationType:
            type_counts[eq_type] = len([e for e in equations if e.equation_type == eq_type])
        
        total_equations = len(equations)
        
        # Create statistics table
        table = doc.add_table(rows=len(EquationType) + 2, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Equation Type'
        header_cells[1].text = 'Count'
        header_cells[2].text = 'Percentage'
        
        # Data rows
        row_idx = 1
        for eq_type in EquationType:
            count = type_counts[eq_type]
            percentage = (count / total_equations * 100) if total_equations > 0 else 0
            
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = eq_type.value.capitalize()
            row_cells[1].text = str(count)
            row_cells[2].text = f"{percentage:.1f}%"
            row_idx += 1
        
        # Total row
        total_cells = table.rows[row_idx].cells
        total_cells[0].text = 'TOTAL'
        total_cells[1].text = str(total_equations)
        total_cells[2].text = '100%'
        
        # Add source file statistics
        doc.add_paragraph().add_run('Source Files:').bold = True
        source_files = set(eq.source_file for eq in equations)
        for source_file in sorted(source_files)[:10]:  # Show first 10
            file_eq_count = len([e for e in equations if e.source_file == source_file])
            doc.add_paragraph(f"  • {Path(source_file).name}: {file_eq_count} equations")
        
        if len(source_files) > 10:
            doc.add_paragraph(f"  ... and {len(source_files) - 10} more files")
        
        doc.add_page_break()
    
    def _add_equation_section(self, doc, eq_type, equations: List[Equation]):
        """Add a section of equations to the document."""
        if isinstance(eq_type, EquationType):
            section_title = eq_type.value.capitalize() + " Equations"
        else:
            section_title = eq_type
        
        heading = doc.add_heading(section_title, 1)
        heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x5A, 0x87)  # Blue color
        
        if not equations:
            doc.add_paragraph("No equations found in this category.")
            return
        
        # Add table of contents for this section
        if len(equations) > 5:
            toc = doc.add_paragraph()
            toc.add_run("Contents: ").bold = True
            for i, eq in enumerate(equations[:10], 1):  # Show first 10 in TOC
                toc.add_run(f"[{i}] ")
            if len(equations) > 10:
                toc.add_run(f"... ({len(equations)} total)")
        
        # Add each equation
        for idx, equation in enumerate(equations, 1):
            # Equation number
            eq_num = doc.add_paragraph()
            eq_num.add_run(f"Equation {idx}: ").bold = True
            
            # Original content
            content_para = doc.add_paragraph()
            content_para.add_run("Original: ").italic = True
            content_para.add_run(equation.content)
            
            # ASCIIMath representation
            ascii_para = doc.add_paragraph()
            ascii_para.add_run("ASCIIMath: ").italic = True
            ascii_para.add_run(equation.asciimath)
            
            # Metadata
            meta_para = doc.add_paragraph(style='Intense Quote')
            meta_text = f"Source: {Path(equation.source_file).name}"
            if equation.page_number:
                meta_text += f" | Page: {equation.page_number}"
            meta_text += f" | Confidence: {equation.confidence:.2f}"
            if equation.tags:
                meta_text += f" | Tags: {', '.join(equation.tags)}"
            meta_para.add_run(meta_text).font.size = Pt(9)
            meta_para.paragraph_format.left_indent = Pt(20)
            
            # Add separator
            if idx < len(equations):
                doc.add_paragraph("─" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing

class EquationAnalyzer:
    """Analyze and categorize equations."""
    
    def __init__(self):
        self.db = EquationDatabase()
    
    def analyze_equations(self, equations: List[Equation] = None) -> Dict:
        """Perform comprehensive analysis of equations."""
        if equations is None:
            equations = self.db.get_all_equations()
        
        analysis = {
            'total_equations': len(equations),
            'by_type': {},
            'by_source': {},
            'by_confidence': {
                'high': 0,    # >= 0.8
                'medium': 0,  # 0.5-0.79
                'low': 0      # < 0.5
            },
            'complexity': {
                'simple': 0,    # < 20 chars
                'medium': 0,    # 20-100 chars
                'complex': 0    # > 100 chars
            },
            'common_patterns': [],
            'recent_additions': []
        }
        
        if not equations:
            return analysis
        
        # Analyze by type
        for eq_type in EquationType:
            type_eqs = [e for e in equations if e.equation_type == eq_type]
            analysis['by_type'][eq_type.value] = {
                'count': len(type_eqs),
                'percentage': len(type_eqs) / len(equations) * 100
            }
        
        # Analyze by source
        source_counts = {}
        for eq in equations:
            source = Path(eq.source_file).name
            source_counts[source] = source_counts.get(source, 0) + 1
        
        analysis['by_source'] = dict(sorted(
            source_counts.items(), 
            key=lambda x: x[1], 
            reverse=True
        )[:10])  # Top 10 sources
        
        # Analyze by confidence
        for eq in equations:
            if eq.confidence >= 0.8:
                analysis['by_confidence']['high'] += 1
            elif eq.confidence >= 0.5:
                analysis['by_confidence']['medium'] += 1
            else:
                analysis['by_confidence']['low'] += 1
        
        # Analyze complexity
        for eq in equations:
            length = len(eq.content)
            if length < 20:
                analysis['complexity']['simple'] += 1
            elif length < 100:
                analysis['complexity']['medium'] += 1
            else:
                analysis['complexity']['complex'] += 1
        
        # Find common patterns
        all_content = ' '.join([eq.content for eq in equations])
        words = re.findall(r'\b[a-zA-Z]+\b', all_content)
        from collections import Counter
        common_words = Counter(words).most_common(10)
        analysis['common_patterns'] = common_words
        
        # Recent additions
        recent = sorted(equations, 
                       key=lambda x: x.extracted_date, 
                       reverse=True)[:5]
        analysis['recent_additions'] = [
            {
                'content': eq.content[:50] + ('...' if len(eq.content) > 50 else ''),
                'type': eq.equation_type.value,
                'date': eq.extracted_date
            }
            for eq in recent
        ]
        
        return analysis
    
    def suggest_categorization(self, equation: Equation) -> List[Tuple[EquationType, float]]:
        """Suggest possible categories for an equation with confidence scores."""
        suggestions = []
        equation_lower = equation.content.lower()
        
        # Check against patterns for each type
        for eq_type in EquationType:
            score = 0
            patterns = EquationExtractor.EQUATION_PATTERNS.get(eq_type, [])
            
            for pattern in patterns:
                try:
                    if re.search(pattern, equation.content, re.IGNORECASE):
                        score += 1
                except:
                    continue
            
            # Check for keywords
            type_keywords = {
                EquationType.PHYSICS: ['force', 'energy', 'velocity', 'acceleration', 'mass'],
                EquationType.CHEMISTRY: ['h2o', 'co2', 'reaction', 'molecule', 'atom'],
                EquationType.BIOLOGY: ['cell', 'dna', 'protein', 'enzyme', 'organism'],
                EquationType.MATHEMATICAL: ['integral', 'derivative', 'function', 'matrix', 'vector']
            }
            
            for keyword in type_keywords.get(eq_type, []):
                if keyword in equation_lower:
                    score += 0.5
            
            if score > 0:
                suggestions.append((eq_type, min(score / 5, 1.0)))
        
        # Sort by confidence score
        suggestions.sort(key=lambda x: x[1], reverse=True)
        
        return suggestions

class InteractiveCLI:
    """Command-line interface for the equation extractor."""
    
    def __init__(self):
        self.extractor = EquationExtractor()
        self.analyzer = EquationAnalyzer()
        self.current_output_file = None
    
    def run(self):
        """Run the interactive CLI."""
        print("=" * 60)
        print("EQUATION EXTRACTOR AND COMPILER")
        print("=" * 60)
        
        while True:
            print("\n" + "=" * 40)
            print("MAIN MENU")
            print("=" * 40)
            print("1. Extract equations from file/directory")
            print("2. Set output Word file location")
            print("3. Compile equations to Word")
            print("4. View extracted equations")
            print("5. Analyze equations")
            print("6. Re-categorize equations")
            print("7. Export equations to JSON")
            print("8. Import equations from JSON")
            print("9. Clear database")
            print("0. Exit")
            print("-" * 40)
            
            choice = input("Enter your choice (0-9): ").strip()
            
            if choice == '1':
                self.extract_equations()
            elif choice == '2':
                self.set_output_file()
            elif choice == '3':
                self.compile_to_word()
            elif choice == '4':
                self.view_equations()
            elif choice == '5':
                self.analyze_equations()
            elif choice == '6':
                self.recategorize_equations()
            elif choice == '7':
                self.export_to_json()
            elif choice == '8':
                self.import_from_json()
            elif choice == '9':
                self.clear_database()
            elif choice == '0':
                print("Goodbye!")
                break
            else:
                print("Invalid choice. Please try again.")
    
    def extract_equations(self):
        """Extract equations from file or directory."""
        print("\n" + "-" * 40)
        print("EXTRACT EQUATIONS")
        print("-" * 40)
        
        path = input("Enter file or directory path: ").strip()
        
        if not path:
            print("No path provided.")
            return
        
        path = Path(path)
        
        if not path.exists():
            print(f"Path does not exist: {path}")
            return
        
        if path.is_file():
            print(f"Processing file: {path.name}")
            equations = self.extractor.extract_from_file(str(path))
        else:
            print(f"Processing directory: {path}")
            equations = self.extractor.extract_from_directory(str(path))
        
        print(f"\nExtracted {len(equations)} equations.")
        
        # Show summary by type
        if equations:
            print("\nExtraction Summary:")
            for eq_type in EquationType:
                count = len([e for e in equations if e.equation_type == eq_type])
                if count > 0:
                    print(f"  {eq_type.value.capitalize()}: {count}")
    
    def set_output_file(self):
        """Set the output Word file location."""
        print("\n" + "-" * 40)
        print("SET OUTPUT FILE")
        print("-" * 40)
        
        path = input("Enter output Word file path (.docx): ").strip()
        
        if not path.endswith('.docx'):
            path += '.docx'
        
        self.current_output_file = path
        print(f"Output file set to: {path}")
    
    def compile_to_word(self):
        """Compile equations to Word document."""
        if not self.current_output_file:
            print("No output file set. Please set output file first.")
            return
        
        print("\n" + "-" * 40)
        print("COMPILE TO WORD")
        print("-" * 40)
        
        # Get compilation options
        append = input("Append to existing file? (y/n): ").lower() == 'y'
        sort_by_type = input("Sort equations by type? (y/n): ").lower() == 'y'
        include_stats = input("Include statistics? (y/n): ").lower() == 'y'
        
        # Create compiler
        compiler = EquationCompiler(self.current_output_file)
        
        # Get equations from database
        equations = self.extractor.db.get_all_equations()
        
        if not equations:
            print("No equations in database to compile.")
            return
        
        print(f"Compiling {len(equations)} equations...")
        
        success = compiler.compile_to_word(
            equations=equations,
            append=append,
            sort_by_type=sort_by_type,
            include_statistics=include_stats
        )
        
        if success:
            print(f"Successfully compiled to {self.current_output_file}")
        else:
            print("Compilation failed.")
    
    def view_equations(self):
        """View extracted equations."""
        print("\n" + "-" * 40)
        print("VIEW EQUATIONS")
        print("-" * 40)
        
        equations = self.extractor.db.get_all_equations()
        
        if not equations:
            print("No equations in database.")
            return
        
        print(f"Total equations: {len(equations)}")
        
        # Filter options
        print("\nFilter options:")
        print("1. View all")
        print("2. By type")
        print("3. By source file")
        print("4. Recent additions")
        
        filter_choice = input("Enter filter choice (1-4): ").strip()
        
        if filter_choice == '1':
            filtered_eqs = equations
        elif filter_choice == '2':
            print("\nEquation types:")
            for i, eq_type in enumerate(EquationType, 1):
                count = len([e for e in equations if e.equation_type == eq_type])
                print(f"{i}. {eq_type.value.capitalize()} ({count})")
            
            type_choice = input("Select type number: ").strip()
            try:
                type_idx = int(type_choice) - 1
                if 0 <= type_idx < len(EquationType):
                    selected_type = list(EquationType)[type_idx]
                    filtered_eqs = [e for e in equations if e.equation_type == selected_type]
                else:
                    print("Invalid type selection.")
                    return
            except ValueError:
                print("Invalid input.")
                return
        elif filter_choice == '3':
            source_files = set(e.source_file for e in equations)
            print("\nSource files:")
            for i, source in enumerate(sorted(source_files)[:20], 1):
                count = len([e for e in equations if e.source_file == source])
                print(f"{i}. {Path(source).name} ({count})")
            
            source_choice = input("Select source number: ").strip()
            try:
                source_idx = int(source_choice) - 1
                selected_source = sorted(source_files)[source_idx]
                filtered_eqs = [e for e in equations if e.source_file == selected_source]
            except (ValueError, IndexError):
                print("Invalid source selection.")
                return
        elif filter_choice == '4':
            n = input("How many recent equations to show? (default 10): ").strip()
            n = int(n) if n.isdigit() else 10
            filtered_eqs = sorted(equations, key=lambda x: x.extracted_date, reverse=True)[:n]
        else:
            print("Invalid choice.")
            return
        
        # Display equations
        print(f"\nShowing {len(filtered_eqs)} equations:")
        print("-" * 60)
        
        for i, eq in enumerate(filtered_eqs, 1):
            print(f"\n[{i}] {eq.equation_type.value.upper()}")
            print(f"Content: {eq.content[:100]}{'...' if len(eq.content) > 100 else ''}")
            print(f"ASCIIMath: {eq.asciimath[:80]}{'...' if len(eq.asciimath) > 80 else ''}")
            print(f"Source: {Path(eq.source_file).name}")
            if eq.page_number:
                print(f"Page: {eq.page_number}")
            print(f"Confidence: {eq.confidence:.2f}")
            print(f"Date: {eq.extracted_date[:10]}")
            print("-" * 40)
            
            if i % 5 == 0 and i < len(filtered_eqs):
                cont = input("Press Enter to continue, 'q' to quit: ")
                if cont.lower() == 'q':
                    break
    
    def analyze_equations(self):
        """Analyze extracted equations."""
        print("\n" + "-" * 40)
        print("ANALYZE EQUATIONS")
        print("-" * 40)
        
        analysis = self.analyzer.analyze_equations()
        
        print(f"Total Equations: {analysis['total_equations']}")
        
        print("\nBy Type:")
        for eq_type, stats in analysis['by_type'].items():
            print(f"  {eq_type.c

        print(f"Total Equations: {analysis['total_equations']}")
        
        print("\nBy Type:")
        for eq_type, stats in analysis['by_type'].items():
            print(f"  {eq_type.capitalize()}: {stats['count']} ({stats['percentage']:.1f}%)")
        
        print("\nBy Confidence Level:")
        conf = analysis['by_confidence']
        print(f"  High (≥0.8): {conf['high']}")
        print(f"  Medium (0.5-0.79): {conf['medium']}")
        print(f"  Low (<0.5): {conf['low']}")
        
        print("\nBy Complexity:")
        comp = analysis['complexity']
        print(f"  Simple (<20 chars): {comp['simple']}")
        print(f"  Medium (20-100 chars): {comp['medium']}")
        print(f"  Complex (>100 chars): {comp['complex']}")
        
        print("\nTop 10 Source Files:")
        for i, (source, count) in enumerate(analysis['by_source'].items(), 1):
            print(f"  {i}. {source}: {count}")
        
        print("\nMost Common Terms:")
        for term, freq in analysis['common_patterns']:
            print(f"  {term}: {freq}")
        
        print("\nRecent Additions:")
        for recent in analysis['recent_additions']:
            print(f"  {recent['date'][:10]}: {recent['content']} ({recent['type']})")
    
    def recategorize_equations(self):
        """Re-categorize equations using machine learning."""
        print("\n" + "-" * 40)
        print("RE-CATEGORIZE EQUATIONS")
        print("-" * 40)
        
        equations = self.extractor.db.get_all_equations()
        
        if len(equations) < 10:
            print("Need at least 10 equations for re-categorization.")
            return
        
        print(f"Found {len(equations)} equations for re-categorization.")
        
        # Use TF-IDF and clustering
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.cluster import KMeans
        
        # Prepare text data
        texts = [eq.content for eq in equations]
        
        # Create TF-IDF features
        vectorizer = TfidfVectorizer(
            max_features=1000,
            stop_words='english',
            ngram_range=(1, 2)
        )
        
        try:
            X = vectorizer.fit_transform(texts)
            
            # Determine optimal number of clusters (4 for our types)
            n_clusters = 4
            
            # Apply KMeans clustering
            kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
            clusters = kmeans.fit_predict(X)
            
            # Map clusters to equation types
            cluster_type_map = {}
            for cluster_id in range(n_clusters):
                cluster_eqs = [equations[i] for i in range(len(equations)) if clusters[i] == cluster_id]
                
                # Find most common type in cluster
                type_counts = {}
                for eq in cluster_eqs:
                    type_counts[eq.equation_type] = type_counts.get(eq.equation_type, 0) + 1
                
                if type_counts:
                    most_common_type = max(type_counts.items(), key=lambda x: x[1])[0]
                    cluster_type_map[cluster_id] = most_common_type
            
            # Re-categorize equations
            re_categorized = 0
            for i, eq in enumerate(equations):
                predicted_type = cluster_type_map.get(clusters[i], eq.equation_type)
                
                if predicted_type != eq.equation_type:
                    eq.equation_type = predicted_type
                    self.extractor.db.save_equation(eq)
                    re_categorized += 1
            
            print(f"Re-categorized {re_categorized} equations.")
            
            # Show cluster distribution
            print("\nCluster Distribution:")
            for cluster_id in range(n_clusters):
                cluster_size = sum(1 for c in clusters if c == cluster_id)
                assigned_type = cluster_type_map.get(cluster_id, "Unknown")
                print(f"  Cluster {cluster_id}: {cluster_size} equations -> {assigned_type.value}")
                
        except Exception as e:
            print(f"Error during re-categorization: {e}")
    
    def export_to_json(self):
        """Export equations to JSON file."""
        print("\n" + "-" * 40)
        print("EXPORT TO JSON")
        print("-" * 40)
        
        output_path = input("Enter output JSON file path: ").strip()
        if not output_path:
            print("No output path provided.")
            return
        
        if not output_path.endswith('.json'):
            output_path += '.json'
        
        equations = self.extractor.db.get_all_equations()
        
        # Convert equations to dictionaries
        equations_data = []
        for eq in equations:
            eq_dict = asdict(eq)
            eq_dict['equation_type'] = eq.equation_type.value
            equations_data.append(eq_dict)
        
        # Export to JSON
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump({
                    'metadata': {
                        'export_date': datetime.now().isoformat(),
                        'total_equations': len(equations),
                        'version': '1.0'
                    },
                    'equations': equations_data
                }, f, indent=2, ensure_ascii=False)
            
            print(f"Successfully exported {len(equations)} equations to {output_path}")
            
        except Exception as e:
            print(f"Error exporting to JSON: {e}")
    
    def import_from_json(self):
        """Import equations from JSON file."""
        print("\n" + "-" * 40)
        print("IMPORT FROM JSON")
        print("-" * 40)
        
        input_path = input("Enter input JSON file path: ").strip()
        if not input_path:
            print("No input path provided.")
            return
        
        if not os.path.exists(input_path):
            print(f"File not found: {input_path}")
            return
        
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if 'equations' not in data:
                print("Invalid JSON format: missing 'equations' key")
                return
            
            imported_count = 0
            skipped_count = 0
            
            for eq_data in data['equations']:
                try:
                    # Convert string equation_type back to Enum
                    eq_data['equation_type'] = EquationType(eq_data['equation_type'])
                    
                    # Create Equation object
                    equation = Equation(**eq_data)
                    
                    # Save to database
                    self.extractor.db.save_equation(equation)
                    imported_count += 1
                    
                except Exception as e:
                    logger.warning(f"Failed to import equation: {e}")
                    skipped_count += 1
            
            print(f"Imported {imported_count} equations, skipped {skipped_count}")
            
        except Exception as e:
            print(f"Error importing from JSON: {e}")
    
    def clear_database(self):
        """Clear all equations from database."""
        print("\n" + "-" * 40)
        print("CLEAR DATABASE")
        print("-" * 40)
        
        confirm = input("Are you sure you want to clear ALL equations? (yes/no): ").strip().lower()
        
        if confirm == 'yes':
            # Get all equations and delete them
            equations = self.extractor.db.get_all_equations()
            
            for eq in equations:
                self.extractor.db.delete_equation(eq.id)
            
            print(f"Cleared {len(equations)} equations from database.")
        else:
            print("Database clear cancelled.")

def main():
    """Main entry point for command-line usage."""
    parser = argparse.ArgumentParser(
        description='Extract equations from various file formats and compile to Word'
    )
    
    parser.add_argument(
        '--input', '-i',
        help='Input file or directory path'
    )
    
    parser.add_argument(
        '--output', '-o',
        help='Output Word file path (.docx)'
    )
    
    parser.add_argument(
        '--append', '-a',
        action='store_true',
        help='Append to existing output file'
    )
    
    parser.add_argument(
        '--interactive', '-I',
        action='store_true',
        help='Launch interactive mode'
    )
    
    parser.add_argument(
        '--export-json',
        help='Export equations to JSON file'
    )
    
    parser.add_argument(
        '--import-json',
        help='Import equations from JSON file'
    )
    
    parser.add_argument(
        '--analyze',
        action='store_true',
        help='Analyze extracted equations'
    )
    
    parser.add_argument(
        '--view',
        action='store_true',
        help='View extracted equations'
    )
    
    args = parser.parse_args()
    
    if args.interactive or not any(vars(args).values()):
        # Launch interactive mode
        cli = InteractiveCLI()
        cli.run()
        return
    
    # Non-interactive mode
    extractor = EquationExtractor()
    
    # Import from JSON if specified
    if args.import_json:
        print(f"Importing equations from {args.import_json}...")
        # Create a temporary CLI for import functionality
        cli = InteractiveCLI()
        cli.extractor = extractor
        cli.import_from_json = lambda: None  # Override for direct call
        # We'll need to implement direct import here
        try:
            with open(args.import_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if 'equations' in data:
                imported = 0
                for eq_data in data['equations']:
                    try:
                        eq_data['equation_type'] = EquationType(eq_data['equation_type'])
                        equation = Equation(**eq_data)
                        extractor.db.save_equation(equation)
                        imported += 1
                    except:
                        pass
                print(f"Imported {imported} equations")
        except Exception as e:
            print(f"Import failed: {e}")
    
    # Extract equations if input specified
    if args.input:
        input_path = Path(args.input)
        
        if input_path.is_file():
            print(f"Processing file: {input_path}")
            equations = extractor.extract_from_file(str(input_path))
        else:
            print(f"Processing directory: {input_path}")
            equations = extractor.extract_from_directory(str(input_path))
        
        print(f"Extracted {len(equations)} equations")
    
    # Analyze if requested
    if args.analyze:
        analyzer = EquationAnalyzer()
        analysis = analyzer.analyze_equations()
        print(f"\nAnalysis:")
        print(f"Total equations: {analysis['total_equations']}")
        for eq_type, stats in analysis['by_type'].items():
            print(f"  {eq_type}: {stats['count']}")
    
    # View if requested
    if args.view:
        equations = extractor.db.get_all_equations()
        print(f"\nTotal equations in database: {len(equations)}")
        if equations:
            print("\nRecent equations:")
            for eq in equations[-5:]:
                print(f"  {eq.equation_type.value}: {eq.content[:50]}...")
    
    # Export to JSON if specified
    if args.export_json:
        equations = extractor.db.get_all_equations()
        equations_data = []
        for eq in equations:
            eq_dict = asdict(eq)
            eq_dict['equation_type'] = eq.equation_type.value
            equations_data.append(eq_dict)
        
        try:
            with open(args.export_json, 'w', encoding='utf-8') as f:
                json.dump({
                    'metadata': {
                        'export_date': datetime.now().isoformat(),
                        'total_equations': len(equations)
                    },
                    'equations': equations_data
                }, f, indent=2, ensure_ascii=False)
            print(f"Exported {len(equations)} equations to {args.export_json}")
        except Exception as e:
            print(f"Export failed: {e}")
    
    # Compile to Word if output specified
    if args.output:
        if not args.output.endswith('.docx'):
            args.output += '.docx'
        
        compiler = EquationCompiler(args.output)
        success = compiler.compile_to_word(append=args.append)
        
        if success:
            print(f"Successfully compiled to {args.output}")
        else:
            print("Compilation failed")

if __name__ == "__main__":
    # Create necessary directories
    os.makedirs('logs', exist_ok=True)
    os.makedirs('exports', exist_ok=True)
    
    # Check for required dependencies
    try:
        import pytesseract
        # Test Tesseract
        pytesseract.get_tesseract_version()
    except:
        print("Warning: Tesseract OCR not properly configured.")
        print("Image OCR functionality may not work.")
    
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nProgram interrupted by user.")
    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        print(f"Fatal error occurred: {e}")
        print("Check equation_extractor.log for details.")


